from __future__ import annotations

import os
import zipfile
import base64
import hashlib
from io import BytesIO
from typing import Tuple, Optional
import tempfile
import subprocess
import threading

from fastapi import FastAPI, UploadFile, File, Header, HTTPException
from pypdf import PdfReader
from docx import Document
import csv
from openpyxl import load_workbook
from PIL import Image, ImageOps, ImageFilter

# Optional OCR dependency
try:
    import pytesseract  # type: ignore
    OCR_AVAILABLE = True
except Exception:
    pytesseract = None
    OCR_AVAILABLE = False

# Optional Whisper dependency
try:
    import whisper  # type: ignore
    WHISPER_AVAILABLE = True
except Exception:
    whisper = None
    WHISPER_AVAILABLE = False

app = FastAPI()

PARSER_SECRET = os.getenv("PARSER_SECRET", "")

MAX_PARSE_BYTES = int(os.getenv("MAX_PARSE_BYTES", "15000000"))  # 15MB
MAX_TEXT_CHARS = int(os.getenv("MAX_TEXT_CHARS", "200000"))      # 200k chars
MAX_PDF_PAGES = int(os.getenv("MAX_PDF_PAGES", "50"))
MAX_DOCX_UNZIPPED_BYTES = int(os.getenv("MAX_DOCX_UNZIPPED_BYTES", "50000000"))  # 50MB

MAX_TABLES = int(os.getenv("MAX_TABLES", "25"))
MAX_ROWS_TOTAL = int(os.getenv("MAX_ROWS_TOTAL", "50000"))
MAX_COLS = int(os.getenv("MAX_COLS", "50"))
MAX_CELL_CHARS = int(os.getenv("MAX_CELL_CHARS", "2000"))

MAX_IMAGES = int(os.getenv("MAX_IMAGES", "50"))
MAX_IMAGE_BYTES = int(os.getenv("MAX_IMAGE_BYTES", "5000000"))  # 5MB each

# Media controls
ENABLE_TRANSCRIBE = os.getenv("ENABLE_TRANSCRIBE", "1") == "1"
MAX_MEDIA_SECONDS = int(os.getenv("MAX_MEDIA_SECONDS", "900"))  # 15 minutes
MAX_AUDIO_BYTES = int(os.getenv("MAX_AUDIO_BYTES", "30000000")) # 30MB
MAX_VIDEO_BYTES = int(os.getenv("MAX_VIDEO_BYTES", "50000000")) # 50MB
WHISPER_MODEL = os.getenv("WHISPER_MODEL", "tiny")              # tiny/base/small...
WHISPER_DEVICE = os.getenv("WHISPER_DEVICE", "cpu")             # cpu/cuda (later)

# Cache whisper model (load once)
_whisper_model = None
_whisper_lock = threading.Lock()


# -------------------------
# Helpers
# -------------------------
def require_secret(x_parser_secret: str) -> None:
    if not PARSER_SECRET or x_parser_secret != PARSER_SECRET:
        raise HTTPException(status_code=401, detail="Unauthorized")


def sanitize_pg_text(s: str) -> str:
    if not s:
        return ""
    s = s.replace("\x00", "")
    s = s.replace("\r\n", "\n").replace("\r", "\n")
    return s


def clamp_text(text: str) -> Tuple[str, bool]:
    if len(text) <= MAX_TEXT_CHARS:
        return text, False
    return text[:MAX_TEXT_CHARS], True


def _clamp_cell(v: object) -> str:
    s = "" if v is None else str(v)
    return s[:MAX_CELL_CHARS] if len(s) > MAX_CELL_CHARS else s


def _b64(b: bytes) -> str:
    return base64.b64encode(b).decode("ascii")


def _safe_name(filename: str) -> str:
    name = (filename or "").replace("\\", "/").split("/")[-1]
    return name or "upload.bin"


def _ensure_ffmpeg() -> None:
    try:
        subprocess.run(["ffmpeg", "-version"], stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL, check=True)
    except Exception:
        raise HTTPException(status_code=500, detail="ffmpeg not available in parser image")


# -------------------------
# Type sniffing
# -------------------------
def sniff_type(filename: str, blob: bytes) -> str:
    name = (filename or "").lower()

    if blob.startswith(b"%PDF-"):
        return "pdf"

    # images
    if blob.startswith(b"\x89PNG\r\n\x1a\n"):
        return "image"
    if blob.startswith(b"\xff\xd8\xff"):
        return "image"
    if blob[:4] == b"RIFF" and blob[8:12] == b"WEBP":
        return "image"
    if blob[:6] in (b"GIF87a", b"GIF89a"):
        return "image"

    # audio/video
    if blob[:4] == b"RIFF" and blob[8:12] == b"WAVE":
        return "audio"
    if blob[:4] == b"OggS":
        return "audio"
    if blob[:4] == b"fLaC":
        return "audio"
    if blob[:4] == b"\x1a\x45\xdf\xa3":  # MKV/WebM
        # webm can be audio OR video; treat as video here
        return "video"
    if len(blob) >= 12 and blob[4:8] == b"ftyp":  # mp4/mov
        return "video"

    # zip-based office docs (DOCX/XLSX) → only classify DOCX here
    if blob[:2] == b"PK":
        try:
            with zipfile.ZipFile(BytesIO(blob)) as z:
                if "[Content_Types].xml" in z.namelist():
                    ct = z.read("[Content_Types].xml")
                    if b"wordprocessingml" in ct or name.endswith(".docx"):
                        return "docx"
        except Exception:
            pass

    # extension fallback
    if name.endswith(".pdf"):
        return "pdf"
    if name.endswith(".docx"):
        return "docx"
    if name.endswith((".png", ".jpg", ".jpeg", ".webp", ".gif")):
        return "image"
    if name.endswith((".mp3", ".wav", ".m4a", ".aac", ".flac", ".ogg")):
        return "audio"
    if name.endswith((".mp4", ".mov", ".mkv", ".webm")):
        return "video"
    if name.endswith((".txt", ".md")):
        return "txt"

    return "txt"


# -------------------------
# Text extraction
# -------------------------
def extract_pdf_text(blob: bytes) -> Tuple[str, dict]:
    reader = PdfReader(BytesIO(blob))
    if getattr(reader, "is_encrypted", False):
        raise HTTPException(status_code=400, detail="Encrypted PDF not supported")

    texts = []
    pages = min(len(reader.pages), MAX_PDF_PAGES)
    for i in range(pages):
        t = reader.pages[i].extract_text() or ""
        if t:
            texts.append(t)
        if sum(len(x) for x in texts) > MAX_TEXT_CHARS * 2:
            break

    joined = sanitize_pg_text("\n".join(texts).strip())
    joined, truncated = clamp_text(joined)
    return joined, {"type": "pdf", "pages_read": pages, "truncated": truncated}


def _docx_unzipped_size(blob: bytes) -> int:
    with zipfile.ZipFile(BytesIO(blob)) as z:
        total = 0
        for info in z.infolist():
            total += int(info.file_size)
            if total > MAX_DOCX_UNZIPPED_BYTES:
                break
        return total


def extract_docx_text(blob: bytes) -> Tuple[str, dict]:
    if _docx_unzipped_size(blob) > MAX_DOCX_UNZIPPED_BYTES:
        raise HTTPException(status_code=413, detail="DOCX unzipped size too large")

    doc = Document(BytesIO(blob))

    parts = []
    for p in doc.paragraphs:
        if p.text:
            parts.append(p.text)
        if sum(len(x) for x in parts) > MAX_TEXT_CHARS * 2:
            break

    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text)
            if row_text:
                parts.append(row_text)
            if sum(len(x) for x in parts) > MAX_TEXT_CHARS * 2:
                break

    joined = sanitize_pg_text("\n".join(parts).strip())
    joined, truncated = clamp_text(joined)
    return joined, {"type": "docx", "truncated": truncated}


def extract_txt(blob: bytes) -> Tuple[str, dict]:
    text = sanitize_pg_text(blob.decode("utf-8", errors="ignore").strip())
    text, truncated = clamp_text(text)
    return text, {"type": "txt", "truncated": truncated}


# -------------------------
# Tables extraction
# -------------------------
def _sniff_table_type(filename: str, blob: bytes) -> str:
    name = (filename or "").lower()
    if name.endswith(".csv"):
        return "csv"
    if name.endswith(".xlsx"):
        return "xlsx"
    if blob[:2] == b"PK":
        return "xlsx"
    return "csv"


def _parse_csv(blob: bytes) -> list[dict]:
    text = blob.decode("utf-8", errors="ignore")
    reader = csv.reader(text.splitlines())
    rows = []
    for row in reader:
        rows.append([_clamp_cell(c) for c in row[:MAX_COLS]])
        if len(rows) >= MAX_ROWS_TOTAL:
            break
    if not rows:
        return []
    header = rows[0]
    data_rows = rows[1:] if len(rows) > 1 else []
    return [{"sheet_name": None, "table_index": 0, "columns": header, "rows": data_rows}]


def _parse_xlsx(blob: bytes) -> list[dict]:
    wb = load_workbook(BytesIO(blob), read_only=True, data_only=True)
    out: list[dict] = []
    total_rows = 0
    table_index = 0

    for ws in wb.worksheets:
        if len(out) >= MAX_TABLES:
            break

        sheet_rows: list[list[str]] = []
        for r in ws.iter_rows(values_only=True):
            row = [_clamp_cell(v) for v in list(r)[:MAX_COLS]]
            if all((c.strip() == "" for c in row)):
                continue
            sheet_rows.append(row)
            total_rows += 1
            if total_rows >= MAX_ROWS_TOTAL:
                break

        if not sheet_rows:
            continue

        header = sheet_rows[0]
        data_rows = sheet_rows[1:] if len(sheet_rows) > 1 else []

        out.append({"sheet_name": ws.title, "table_index": table_index, "columns": header, "rows": data_rows})
        table_index += 1
        if total_rows >= MAX_ROWS_TOTAL:
            break

    return out


# -------------------------
# Images extraction
# -------------------------
def _image_meta(img_bytes: bytes) -> dict:
    sha = hashlib.sha256(img_bytes).hexdigest()
    width = height = None
    fmt = None
    try:
        with Image.open(BytesIO(img_bytes)) as im:
            width, height = im.size
            fmt = (im.format or "").upper() if im.format else None
    except Exception:
        pass
    return {"sha256": sha, "bytes": len(img_bytes), "width": width, "height": height, "format": fmt}


def _extract_images_docx(blob: bytes) -> list[dict]:
    imgs: list[dict] = []
    with zipfile.ZipFile(BytesIO(blob)) as z:
        media_files = [n for n in z.namelist() if n.startswith("word/media/")]
        for name in media_files[:MAX_IMAGES]:
            data = z.read(name)
            if len(data) > MAX_IMAGE_BYTES:
                continue
            imgs.append({"index": len(imgs), "name": name.split("/")[-1], "meta": _image_meta(data), "b64": _b64(data)})
    return imgs


def _extract_images_pdf(blob: bytes) -> list[dict]:
    imgs: list[dict] = []
    reader = PdfReader(BytesIO(blob))
    if getattr(reader, "is_encrypted", False):
        return []

    for page_index, page in enumerate(reader.pages[:MAX_PDF_PAGES]):
        try:
            page_images = getattr(page, "images", None)
            if not page_images:
                continue
            for im in page_images:
                data = getattr(im, "data", None)
                if not data:
                    continue
                data = bytes(data)
                if len(data) > MAX_IMAGE_BYTES:
                    continue
                imgs.append(
                    {"index": len(imgs), "name": f"page{page_index}_img{len(imgs)}", "meta": _image_meta(data), "b64": _b64(data)}
                )
                if len(imgs) >= MAX_IMAGES:
                    return imgs
        except Exception:
            continue
    return imgs


# -------------------------
# OCR (stronger)
# -------------------------
def _preprocess_variants(im: Image.Image) -> list[Image.Image]:
    # Ensure we’re in RGB then grayscale (handles RGBA/LA/P)
    if im.mode not in ("RGB", "L"):
        im = im.convert("RGB")

    variants: list[Image.Image] = []

    # Base grayscale + upscale if small
    g = ImageOps.grayscale(im)
    w, h = g.size
    scale = 2 if max(w, h) < 1800 else 1
    if scale != 1:
        g = g.resize((w * scale, h * scale))

    g = ImageOps.autocontrast(g)
    variants.append(g)

    # Sharpen
    variants.append(g.filter(ImageFilter.UnsharpMask(radius=2, percent=180, threshold=2)))

    # Binarize (good for screenshots/docs)
    thr = 180
    bw = g.point(lambda p: 255 if p > thr else 0).convert("L")
    variants.append(bw)

    # Invert binarize (good for white text on dark background)
    variants.append(ImageOps.invert(bw))

    return variants


def _run_ocr_best(im: Image.Image) -> Tuple[str, dict]:
    if not OCR_AVAILABLE or pytesseract is None:
        return "", {"ocr_available": False, "engine": "none"}

    configs = [
        "--oem 3 --psm 6",
        "--oem 3 --psm 11",
        "--oem 1 --psm 6",
        "--oem 1 --psm 11",
    ]

    best_text = ""
    best_meta = {"ocr_available": True, "engine": "tesseract", "best_cfg": None, "variant": None}

    variants = _preprocess_variants(im)
    for vi, v in enumerate(variants):
        for cfg in configs:
            try:
                txt = pytesseract.image_to_string(v, config=cfg) or ""
                txt = txt.strip()
                if len(txt) > len(best_text):
                    best_text = txt
                    best_meta["best_cfg"] = cfg
                    best_meta["variant"] = vi
            except Exception:
                continue

    return best_text, best_meta


# -------------------------
# Transcription (ffmpeg + whisper)
# -------------------------
def _load_whisper_model():
    global _whisper_model
    if not WHISPER_AVAILABLE or whisper is None:
        raise HTTPException(status_code=501, detail="Transcription not available (install whisper)")

    with _whisper_lock:
        if _whisper_model is None:
            _whisper_model = whisper.load_model(WHISPER_MODEL, device=WHISPER_DEVICE)
    return _whisper_model


def _ffprobe_duration_seconds(path: str) -> Optional[float]:
    try:
        p = subprocess.run(
            ["ffprobe", "-v", "error", "-show_entries", "format=duration", "-of", "default=nw=1:nk=1", path],
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            text=True,
            timeout=20,
        )
        if p.returncode != 0:
            return None
        s = (p.stdout or "").strip()
        return float(s) if s else None
    except Exception:
        return None


def _run_ffmpeg_extract_audio(src_path: str, out_wav_path: str) -> None:
    _ensure_ffmpeg()
    cmd = [
        "ffmpeg", "-y",
        "-i", src_path,
        "-vn",
        "-ac", "1",
        "-ar", "16000",
        "-t", str(MAX_MEDIA_SECONDS),
        "-f", "wav",
        out_wav_path,
    ]
    p = subprocess.run(cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True)
    if p.returncode != 0:
        raise HTTPException(status_code=500, detail=f"ffmpeg failed: {p.stderr[-800:]}")


def _whisper_transcribe(wav_path: str) -> str:
    if not ENABLE_TRANSCRIBE:
        return ""
    model = _load_whisper_model()
    res = model.transcribe(wav_path)
    return (res.get("text") or "").strip()


def _transcribe_blob(filename: str, blob: bytes) -> Tuple[str, dict]:
    if not ENABLE_TRANSCRIBE:
        return "", {"type": "transcribe", "enabled": False}

    ftype = sniff_type(filename, blob)
    if ftype == "audio" and len(blob) > MAX_AUDIO_BYTES:
        raise HTTPException(status_code=413, detail="Audio too large for transcription")
    if ftype == "video" and len(blob) > MAX_VIDEO_BYTES:
        raise HTTPException(status_code=413, detail="Video too large for transcription")

    safe_name = _safe_name(filename)

    with tempfile.TemporaryDirectory() as td:
        src_path = os.path.join(td, safe_name)
        with open(src_path, "wb") as f:
            f.write(blob)

        dur = _ffprobe_duration_seconds(src_path)

        wav_path = os.path.join(td, "audio.wav")
        _run_ffmpeg_extract_audio(src_path, wav_path)

        text = _whisper_transcribe(wav_path)
        text = sanitize_pg_text(text)
        text, truncated = clamp_text(text)

        return text, {
            "type": "transcribe",
            "engine": "whisper",
            "model": WHISPER_MODEL,
            "device": WHISPER_DEVICE,
            "duration_seconds": dur,
            "truncated": truncated,
            "enabled": True,
        }


# -------------------------
# Routes
# -------------------------
@app.get("/health")
def health():
    return {
        "ok": True,
        "ocr": OCR_AVAILABLE,
        "whisper": WHISPER_AVAILABLE,
        "enable_transcribe": ENABLE_TRANSCRIBE,
        "whisper_model": WHISPER_MODEL,
        "whisper_device": WHISPER_DEVICE,
    }


@app.post("/parse/text")
async def parse_text(
    file: UploadFile = File(...),
    x_parser_secret: str = Header(default=""),
):
    require_secret(x_parser_secret)

    blob = await file.read()
    if len(blob) > MAX_PARSE_BYTES:
        raise HTTPException(status_code=413, detail="File too large for parser")

    filename = file.filename or "upload.bin"
    ftype = sniff_type(filename, blob)

    if ftype == "pdf":
        text, meta = extract_pdf_text(blob)
    elif ftype == "docx":
        text, meta = extract_docx_text(blob)
    elif ftype == "image":
        # IMPORTANT: /parse/text is NOT OCR. OCR is /parse/ocr.
        text, meta = f"[Image file: {filename}]", {"type": "image", "truncated": False}
    elif ftype in ("audio", "video"):
        text, tmeta = _transcribe_blob(filename, blob)
        if not text.strip():
            text = f"[{ftype.upper()} file: {filename}] (transcription produced no text)"
        meta = {"type": ftype, **tmeta}
    else:
        text, meta = extract_txt(blob)

    return {"text": text, "meta": {**meta, "bytes": len(blob), "filename": filename}}


@app.post("/parse/tables")
async def parse_tables(
    file: UploadFile = File(...),
    x_parser_secret: str = Header(default=""),
):
    require_secret(x_parser_secret)

    blob = await file.read()
    if len(blob) > MAX_PARSE_BYTES:
        raise HTTPException(status_code=413, detail="File too large for parser")

    ftype = _sniff_table_type(file.filename or "", blob)
    tables = _parse_xlsx(blob) if ftype == "xlsx" else _parse_csv(blob)
    tables = tables[:MAX_TABLES]

    return {
        "tables": tables,
        "meta": {
            "type": ftype,
            "bytes": len(blob),
            "filename": file.filename,
            "tables_count": len(tables),
            "max_rows_total": MAX_ROWS_TOTAL,
            "max_cols": MAX_COLS,
        },
    }


@app.post("/parse/images")
async def parse_images(
    file: UploadFile = File(...),
    x_parser_secret: str = Header(default=""),
):
    require_secret(x_parser_secret)

    blob = await file.read()
    if len(blob) > MAX_PARSE_BYTES:
        raise HTTPException(status_code=413, detail="File too large for parser")

    filename = file.filename or "upload.bin"
    ftype = sniff_type(filename, blob)

    if ftype == "docx":
        images = _extract_images_docx(blob)
    elif ftype == "pdf":
        images = _extract_images_pdf(blob)
    elif ftype == "image":
        images = [{"index": 0, "name": filename, "meta": _image_meta(blob), "b64": _b64(blob)}]
    else:
        images = []

    images = images[:MAX_IMAGES]
    return {"images": images, "meta": {"type": ftype, "bytes": len(blob), "filename": filename, "images_count": len(images)}}


@app.post("/parse/ocr")
async def parse_ocr(
    file: UploadFile = File(...),
    x_parser_secret: str = Header(default=""),
):
    require_secret(x_parser_secret)

    blob = await file.read()
    if len(blob) > MAX_PARSE_BYTES:
        raise HTTPException(status_code=413, detail="File too large for parser")

    filename = file.filename or "upload.bin"

    try:
        im = Image.open(BytesIO(blob))
    except Exception:
        return {
            "text": "",
            "meta": {"type": "ocr", "engine": "tesseract", "note": "not an image", "bytes": len(blob), "filename": filename},
        }

    text, ocr_meta = _run_ocr_best(im)
    text = sanitize_pg_text(text)
    text, truncated = clamp_text(text)

    # IMPORTANT: tell you WHY it was empty
    note = None
    if not OCR_AVAILABLE:
        note = "OCR unavailable: pytesseract not installed or failed import"
    elif not text.strip():
        note = "OCR ran but produced empty text (image may have no readable text / too blurry / too small)"

    return {
        "text": text,
        "meta": {
            "type": "ocr",
            "bytes": len(blob),
            "filename": filename,
            "truncated": truncated,
            "ocr_available": OCR_AVAILABLE,
            "note": note,
            **ocr_meta,
        },
    }


@app.post("/parse/transcribe")
async def parse_transcribe(
    file: UploadFile = File(...),
    x_parser_secret: str = Header(default=""),
):
    require_secret(x_parser_secret)

    blob = await file.read()
    if len(blob) > MAX_PARSE_BYTES:
        raise HTTPException(status_code=413, detail="File too large for parser")

    filename = file.filename or "upload.bin"
    ftype = sniff_type(filename, blob)
    if ftype not in ("audio", "video"):
        return {"text": "", "meta": {"type": "transcribe", "note": "not audio/video"}}

    text, meta = _transcribe_blob(filename, blob)
    return {"text": text, "meta": {**meta, "filename": filename, "bytes": len(blob), "whisper_available": WHISPER_AVAILABLE}}
